import re
from docx import Document
from datetime import datetime, timedelta, date


# 100% borrowed from: https://newbedev.com/how-to-use-python-docx-to-replace-text-in-a-word-document-and-save
def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def get_date(week_start, add_days=0, date_format="%d/%m"):
    the_new_date = week_start + timedelta(days=add_days)

    return the_new_date.strftime(date_format)


def get_start_of_week(todays_date, date_format="%d/%m"):
    start_of_week = todays_date - timedelta(days=todays_date.weekday())

    return start_of_week


todays_date_placeholder = re.compile(r"<DATE>")
name_placeholder = re.compile(r"<NAME>")
monday_placeholder = re.compile(r"<MON>")
tuesday_placeholder = re.compile(r"<TUE>")
wednesday_placeholder = re.compile(r"<WED>")
thursday_placeholder = re.compile(r"<THU>")
friday_placeholder = re.compile(r"<FRI>")
days_worked_placeholder = re.compile(r"<DAYSWORKED>")
computercenter_contact_name_placeholder = re.compile(r"<COMPUTACENTERNAME>")


def create_timesheet(name_to_use, computercenter_contact_name):
    name = name_to_use

    # Used to document the date the timesheet is created
    todays_date = datetime.today().strftime("%d/%m/%Y")

    # Assumption that you do timesheets on the week you need them
    week_start = get_start_of_week(datetime.today())

    doc = Document("static/files/timesheet_template.docx")

    docx_replace_regex(doc, name_placeholder, name)
    docx_replace_regex(
        doc, computercenter_contact_name_placeholder, computercenter_contact_name
    )
    docx_replace_regex(doc, todays_date_placeholder, todays_date)
    docx_replace_regex(doc, monday_placeholder, get_date(week_start))
    docx_replace_regex(doc, tuesday_placeholder, get_date(week_start, 1))
    docx_replace_regex(doc, wednesday_placeholder, get_date(week_start, 2))
    docx_replace_regex(doc, thursday_placeholder, get_date(week_start, 3))
    docx_replace_regex(doc, friday_placeholder, get_date(week_start, 4))
    docx_replace_regex(doc, days_worked_placeholder, "5")

    # the path to save seems relative to where i ran it, not the script
    file_path = "{0}-Timesheet-WC-{1}.docx".format(
        name.replace(" ", "-"), week_start.strftime("%d-%m-%Y")
    )
    doc.save(file_path)

    return file_path
